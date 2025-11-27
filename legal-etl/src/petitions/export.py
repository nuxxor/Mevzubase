from __future__ import annotations

from pathlib import Path
from typing import Optional

from .schema import PetitionOutput


def export_docx(output_path: Path, petition: PetitionOutput) -> Path:
    """
    Write petition output to a simple DOCX (headings + paragraphs).
    Requires python-docx.
    """
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "DOCX çıktısı için python-docx kurulu olmalı: pip install python-docx"
        ) from exc

    doc = Document()
    doc.add_heading("Dilekçe Taslağı", level=1)
    for line in petition.text.splitlines():
        if not line.strip():
            doc.add_paragraph("")
            continue
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def export_pdf(output_path: Path, petition: PetitionOutput, html: Optional[str] = None) -> Path:
    """
    Export to PDF if WeasyPrint is available; otherwise raise a clear error.
    """
    try:
        from weasyprint import HTML  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "PDF çıktısı için weasyprint kurulu olmalı: pip install weasyprint"
        ) from exc

    html_content = html or petition.html or petition.text.replace("\n", "<br/>")
    HTML(string=html_content).write_pdf(output_path)
    return output_path
